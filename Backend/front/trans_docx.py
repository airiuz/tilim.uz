from .translit_text import to_cyrillic, to_latin
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# def transliterate(file_path, file_path2, t):
#     mytranslate = lambda text:to_cyrillic(text) if t == '1' else to_latin(text)
#     doc_obj = Document(file_path)
    
#     for p in doc_obj.paragraphs:
#         if p.text and p.text.strip():
#             p.text = mytranslate(p.text)
        
#     for table in doc_obj.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 if cell.text and cell.text.strip():
#                     matn = mytranslate(cell.text)
#                     cell.text = matn
#     doc_obj.save(file_path2)


import re
import shutil
from zipfile import ZipFile, ZIP_DEFLATED


def transliterate(file_path1, file_path2, t):
    mytransliterate = lambda text: to_cyrillic(text) if t == '1' else to_latin(text) 
    def convert(match_obj):
            data = match_obj.groupdict()
            return b"<w:t " + data['attrs'] + b">" + mytransliterate(data['text'].decode("utf-8")).encode() + b"</w:t>"

    with ZipFile(file_path1, 'r') as inzip, ZipFile(file_path2, 'w', ZIP_DEFLATED) as outzip:
        for inzipinfo in inzip.infolist():
            with inzip.open(inzipinfo) as infile:
                if inzipinfo.filename == "word/document.xml":
                    new_content = re.sub(rb"<w:t(?P<attrs>[^>]*)>(?P<text>[^<]+)<\/w:t>", convert, infile.read())
                    outzip.writestr(inzipinfo.filename, new_content)
                    # print(new_content)
                else:
                    outzip.writestr(inzipinfo.filename, infile.read())

